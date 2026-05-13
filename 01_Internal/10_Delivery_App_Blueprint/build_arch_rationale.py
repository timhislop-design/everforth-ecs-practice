"""
build_arch_rationale.py — ECS Architectural Rationale & Decision Log
Rebuilt with canonical ECS Federal branding via ecs_template.py
Challenge boxes rendered as navy-bordered callouts (teal accent side-box)
"""
import sys, os
HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.abspath(os.path.join(HERE, "..", ".."))
sys.path.insert(0, os.path.join(REPO, "03_Shared", "00_Templates_and_Branding"))
from ecs_template import EcsDocument, DocMeta, Brand, PPR_ORDER, RPR_ORDER, insert_in_order
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

LOGO = os.path.join(REPO, "00_Master_Blueprint", "assets", "everforth_logo.png")
OUT  = os.path.join(HERE, "ECS_ArchRationale_ForShawnReview_INTERNAL.docx")

doc = EcsDocument(
    meta=DocMeta(
        eyebrow="INTERNAL · ARCHITECTURAL REVIEW — FOR SHAWN",
        title="Architectural Rationale\n& Decision Log",
        subtitle="ECS Delivery Intelligence Platform — How we got here, for architectural challenge and peer review",
        audience="Director of Technical Services (Shawn), Senior Director",
        companion_to="Blueprint A · Blueprint B · Collateral Index & Team Review Guide",
        doc_id="INT-DA-ARCH",
        version="1.0",
        status="For Architecture Review",
        confidentiality="Internal Use Only · Confidential",
        running_header_label="Internal · Arch Rationale & Decision Log",
    ),
    logo_path=LOGO,
)

doc.add_cover_page()
doc.page_break()

# How-to-use callout
doc.callout(
    "How to use this document: This is not a proposal — it is a reasoning trace. Every architectural decision "
    "made in Blueprints A and B is explained here: what we considered, what we ruled out, and why we landed "
    "where we did. The Challenge boxes are the questions we know Shawn should push back on. If a better "
    "answer exists, this is the document to mark up."
)

def challenge(text: str):
    """Render a Challenge box: teal-left-border callout with 'CHALLENGE' label."""
    p = doc.doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4' if side != 'left' else '12')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), Brand.TEAL_BRIGHT_HEX)
        pBdr.append(b)
    insert_in_order(pPr, pBdr, PPR_ORDER)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0FAFA')
    insert_in_order(pPr, shd, PPR_ORDER)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    label = p.add_run("CHALLENGE  ")
    label.font.name = "Calibri"; label.font.size = Pt(9); label.font.bold = True
    label.font.color.rgb = Brand.TEAL_DEEP
    rPr = label._r.get_or_add_rPr()
    sp_e = OxmlElement('w:spacing'); sp_e.set(qn('w:val'), '30')
    insert_in_order(rPr, sp_e, RPR_ORDER)
    body = p.add_run(text)
    body.font.name = "Calibri"; body.font.size = Pt(10)
    body.font.color.rgb = Brand.BODY
    return p

# ── 1. STARTING POINT ───────────────────────────────────────────────────────
doc.h1("Starting Point — What Problem Are We Actually Solving?")
doc.para(
    "The ECS practice has built a comprehensive delivery methodology: consultant handbooks, accelerator packs, "
    "facilitator guides, decision guides, workshop pre-reads, sprint workbooks. The content is solid. The "
    "problem is that it lives in files — Word docs, Excel sheets, PowerPoints — and those files are manually "
    "maintained, quickly stale, and disconnected from the actual work happening in ServiceNow during delivery."
)
doc.para("More specifically, the pain points that drove this initiative were:")
doc.bullet("Delivery managers spending 2–4 hours per week maintaining RAG status spreadsheets that are immediately out of date.")
doc.bullet("Risk logs and decision trackers living in Excel, manually updated, with no notification when something goes overdue.")
doc.bullet("Customers receiving weekly status emails rather than having a live view of their engagement.")
doc.bullet("New consultants needing to read five documents to understand what to do in a given sprint.")
doc.bullet("No cross-engagement visibility for practice leadership — no way to see portfolio health without asking each DM.")
doc.para(
    "The question we started with was: can we fix this with a tool that lives inside ServiceNow itself, so the "
    "delivery intelligence is tied directly to the work, not maintained separately?"
)
challenge(
    "Is the pain real enough to justify the investment? Have we quantified the weekly overhead cost of manual "
    "status maintenance per DM, and compared it honestly to the build investment? If the average DM saves 2 "
    "hours/week across an 18-week engagement, that is 36 hours per engagement. At what engagement volume does "
    "the ROI turn positive, and are we there yet?"
)

# ── 2. WHY A SERVICENOW APPLICATION ─────────────────────────────────────────
doc.h1("Why a ServiceNow Application — Not a Standalone Tool")
doc.para("The first major decision was where the platform lives. Three options were considered:")

doc.h2("Option 1: Hosted on Everforth's ServiceNow Instance (Rejected)")
doc.para(
    "Running a shared delivery portal on Everforth's own instance with customer portal access was the initial "
    "instinct. It was rejected for three reasons. First, ServiceNow's licensing model is not designed for "
    "multi-tenancy — isolating customer A's data from customer B's requires ACL engineering that becomes fragile "
    "at scale. Second, enterprise customers will not accept their project data living on a vendor's instance "
    "they do not control. Third, it puts Everforth in the business of running infrastructure, which is not "
    "the model."
)
challenge(
    "Is multi-tenancy actually as hard as we think? Some partners do run shared SNow instances for managed "
    "services. Are there scoping patterns (domain separation, application scoping) that make this viable? "
    "We dismissed it quickly — Shawn should confirm whether this was the right call."
)

doc.h2("Option 2: Standalone Web Application (Rejected)")
doc.para(
    "A standalone app (custom web app, or something like Notion/Confluence with role-based views) was considered "
    "and rejected. The primary reason: it breaks the meta-narrative. We are an OOTB ServiceNow practice. If we "
    "cannot manage our own delivery using ServiceNow, we have a credibility problem in every sales conversation. "
    "A standalone app also requires separate authentication, hosting, maintenance, and data synchronization — "
    "none of which we want to own."
)
challenge(
    "Is the meta-narrative argument actually compelling to customers, or is it something we tell ourselves? "
    "Has anyone in a real sales conversation asked 'do you use ServiceNow to manage your own delivery?' If not, "
    "we may be solving an internal story problem, not a customer one."
)

doc.h2("Option 3: App on Customer Instance (Selected)")
doc.para(
    "Installing the app on the customer's instance per engagement was selected because: it uses the customer's "
    "existing licensing, keeps their data on their infrastructure, demonstrates OOTB delivery methodology in "
    "practice, and is architecturally clean. The trade-off is that updates to the app must be pushed "
    "per-instance — there is no central update mechanism. We judged this acceptable for the near term."
)
challenge(
    "What is the upgrade management model at scale? If we have 20 customers on the app and we push a v1.1 "
    "update, what does that deployment process look like operationally? Is this a managed service commitment, "
    "a self-service update set, or something else? We have not fully designed the post-launch operations model."
)

# ── 3. WHY NOT NATIVE MODULES AS-IS ─────────────────────────────────────────
doc.h1("Why Not Just Use Native ServiceNow Modules As-Is?")
doc.para(
    "A fair challenge: SPM already has projects and milestones. Agile already has sprints and stories. "
    "Knowledge already has articles. Why build anything new at all — why not just configure what already exists "
    "and point a portal at it?"
)
doc.para(
    "The answer is that native modules do not carry delivery methodology context. A sprint in Agile is a "
    "container for stories — it knows nothing about ECS gates, health scoring, OOTB decision tracking, or "
    "methodology content by role. To surface the right thing to the right person at the right sprint requires "
    "a layer that understands the delivery framework. That layer is what both blueprints provide."
)
doc.para(
    "Both Blueprint A and Blueprint B read from native modules — they do not replicate them. The difference is "
    "that Blueprint A adds a purpose-built data model on top, while Blueprint B extends the native model with "
    "delivery-specific fields. Neither approach replaces Agile, SPM, or the Knowledge Base."
)
challenge(
    "Could we achieve 80% of the value with a well-configured Service Portal sitting directly on native tables, "
    "with zero extension or custom tables? No new data model — just smart widget queries and role-filtered views. "
    "We assumed we need a delivery data model. Is that assumption correct, or is it scope creep dressed as architecture?"
)

# ── 4. THE DATA MODEL ────────────────────────────────────────────────────────
doc.h1("The Data Model — Why These Tables?")
doc.h2("Blueprint A Custom Tables")
doc.table(
    headers=["Table", "Why Custom?", "The Alternative Considered", "Why We Rejected It"],
    rows=[
        ["Engagement", "One master record per customer engagement ties everything together. No native SNow table represents 'an implementation engagement' — pm_project is closest but carries too much SPM baggage and is tied to licensing.", "Extend pm_project", "Decided against for Blueprint A: cleaner to own the record entirely, avoid SPM licensing dependency for the core record, and align to Blueprint A's clean data model philosophy."],
        ["Sprint Gate", "Gates are different from tasks or milestones. They have specific delivery criteria: velocity %, open risk count, decision count, sign-off. No native table models this combination.", "Extend pm_project_task", "Blueprint B uses exactly this approach. Blueprint A avoids it to maintain schema cleanliness and avoid upgrade risk on project_task."],
        ["Risk/Issue", "Delivery risks are different from GRC risks. We need type (risk vs. issue), engagement context, sprint gate reference, and health impact — fields that do not exist on sn_risk_risk without extensions.", "Extend sn_risk_risk", "Blueprint B extends sn_risk_risk. Blueprint A owns the table to avoid GRC licensing dependency and keep the data model self-contained."],
        ["Decision Record", "Decisions are a first-class delivery artifact — they need workstream tagging, impact-if-deferred, and tight coupling to sprint gates. No native table carries this.", "Extend sc_task", "Blueprint B extends sc_task. The catalog task structure is workable but adds unnecessary catalog overhead for what is essentially a governance record."],
        ["Methodology Content", "The content table tags articles by role, sprint, workstream, and audience for portal surfacing. KB articles could be extended, but KB has its own lifecycle and approval process that conflicts with delivery content management.", "Extend kb_knowledge", "Blueprint B extends kb_knowledge with flag fields. Blueprint A owns content to avoid KB governance conflicts and enable richer metadata without affecting the KB workflow."],
    ],
    col_widths_in=[1.5, 2.5, 1.8, 3.56],
)
challenge(
    "Is the Methodology Content table the right call? If we put delivery methodology content into a custom table, "
    "we are building a second knowledge base. Every time a consultant searches KB, they search in two places. "
    "KB has built-in taxonomy, search, and lifecycle. The Blueprint B approach of extending KB with u_ecs_ tags "
    "may be the right answer even in a Blueprint A world. This is the one table worth reconsidering."
)

# ── 5. THE HEALTH SCORE ──────────────────────────────────────────────────────
doc.h1("The Health Score — Why This Formula?")
doc.para(
    "The health score (0–100, driving Red/Yellow/Green) was designed to replace the subjective weekly RAG status "
    "that a delivery manager currently sets manually. The formula was designed with four components because these "
    "are the four things that actually determine whether an engagement is on track:"
)
doc.table(
    headers=["Component", "Weight", "Reasoning", "The Challenge"],
    rows=[
        ["Sprint Velocity", "40%", "If stories are not getting completed, the engagement is not delivering. Velocity is the most direct signal of delivery health and the hardest to mask.", "40% assumes story completion is the primary health signal. On some engagements — particularly early phases where content is design-heavy — low story count does not mean poor health. The formula may unfairly penalize Sprint 1 and 2."],
        ["Risk Posture", "30%", "Open high risks are the leading indicator of future problems. Weighting risks at 30% forces risk management to be visible, not an afterthought.", "The deduction formula (−15 per high risk, −5 per medium) was estimated, not empirically derived. After 3–4 real engagements, these weights may need recalibration."],
        ["Decision Latency", "20%", "Deferred customer decisions are the most common cause of engagement delays. Making decision latency visible in the health score creates urgency.", "A decision 5 days overdue gets the same treatment whether it is about icon colors on the portal or CSDM data model scope. Severity weighting by workstream impact is missing."],
        ["Milestone Variance", "10%", "SPM milestone dates represent the contractual timeline. Variance here is the most obvious signal for a sponsor.", "10% feels low if a milestone is 3 weeks late. Should milestone variance have a non-linear penalty — small variance ignored, large variance heavily weighted?"],
    ],
    col_widths_in=[1.6, 0.7, 2.8, 4.26],
)
challenge(
    "The formula was designed by a delivery manager and a consultant, not by analyzing historical engagement data. "
    "We do not have empirical evidence that these weights predict engagement outcomes. Shawn: is there a better "
    "way to design this, or a dataset we can use to validate it before we build it in?"
)

# ── 6. PORTAL DESIGN RATIONALE ───────────────────────────────────────────────
doc.h1("Portal Design Rationale")
doc.para(
    "The role-based portal was designed around a single principle: every user should see exactly what they need "
    "and nothing they do not. The role/view mapping was derived from how we currently manage communications on "
    "live engagements:"
)
doc.bullet("Sponsors ask: are we on track, when is go-live, what decisions do we need from me? → Sponsor view surfaces these three things and nothing else.")
doc.bullet("Process owners ask: what is coming in my process area, what do I need to read before the workshop? → Process owner view surfaces workstream-scoped content and upcoming workshops.")
doc.bullet("Delivery managers ask: what is the overall health, where are the risks, what is overdue? → DM dashboard surfaces calculated health, risk log, and decision tracker.")
doc.bullet("Consultants ask: what am I doing this sprint, where is the guide for this workstream? → Consultant landing page surfaces their Agile stories and relevant methodology content.")
challenge(
    "We designed the portal based on what we think these roles ask. We have not done a single user interview "
    "with a real customer sponsor or process owner. The portal experience should be validated with at least one "
    "customer before Phase 2 is locked. Is there a willing customer we can show wireframes to before we build?"
)

# ── 7. BUILD VS. BUY ─────────────────────────────────────────────────────────
doc.h1("Build vs. Buy — Did We Ask the Right Question?")
doc.para(
    "At no point in this process did we formally evaluate whether an existing product already solves this "
    "problem. Before committing to build, Shawn should pressure-test the following alternatives:"
)
doc.table(
    headers=["Alternative", "What It Offers", "Why We Likely Still Build"],
    rows=[
        ["ServiceNow Now Assist / GenAI Studio", "Guided delivery content via AI; some health monitoring in ITSM Pro", "Not a delivery management platform — no engagement lifecycle, gates, or customer transparency layer"],
        ["Existing SNow ISV apps on the Store", "Several project delivery tracker apps exist", "None are specific to ECS OOTB methodology delivery; generic PM apps miss the consultant/customer role split and methodology content integration"],
        ["SimpleNow, Fruition, Rego partners", "Partner-built SNow delivery accelerators", "Would need significant customization to match ECS methodology specificity; we would lose IP ownership and the commercial upside"],
        ["Asana / Monday.com with SNow integration", "Modern PM tools with good portal UX", "Adds a second platform; breaks the OOTB meta-narrative; no VA integration; not billable as a SNow deliverable"],
    ],
    col_widths_in=[2.3, 2.8, 4.26],
)
challenge(
    "Has anyone done a thorough ServiceNow Store search for delivery management apps in the last 30 days? The "
    "Store catalogue changes frequently. If there is an app at $5K/year that does 70% of what we need, the "
    "build math changes completely. Shawn: assign someone to do a proper Store competitive scan before Phase 0 kicks off."
)

# ── 8. BLUEPRINT A vs B META-QUESTION ────────────────────────────────────────
doc.h1("Blueprint A vs. Blueprint B — The Meta-Question")
doc.para(
    "Both blueprints were created because the architecture decision has legitimate trade-offs and the right "
    "answer depends on facts we do not yet know. Here is the honest framing of the decision:"
)
doc.h2("Blueprint A is the right choice if:")
doc.bullet("We are committed to the ServiceNow Store as a revenue channel and willing to invest in ISV program enrollment now.")
doc.bullet("We have a dedicated developer available for 16 weeks who will not be pulled onto billable delivery.")
doc.bullet("The App Engine SKU conversation is manageable — either our target customers already have it or it becomes part of the ECS engagement scope.")
doc.bullet("We want a clean commercial product that can be sold, licensed, and supported independent of a consulting engagement.")
doc.h2("Blueprint B is the right choice if:")
doc.bullet("We want a working platform in the hands of consultants and customers by September 2026 rather than October–November.")
doc.bullet("Our immediate goal is proving the concept and replacing the spreadsheets, not building a commercial product.")
doc.bullet("Dedicated Studio app developer availability is uncertain or App Engine SKU availability is uncertain.")
doc.bullet("We are willing to treat Blueprint B as a paid proof-of-concept and accept that some rework may follow if we scale to Blueprint A.")
doc.callout(
    "The honest answer: Blueprint A is the better architecture. Blueprint B is the better first step. The "
    "strongest argument for Blueprint B is not that it is architecturally superior — it is not — but that it "
    "gets the platform into real use faster, generates real feedback, and validates the investment before we "
    "commit the larger sum. If Blueprint B runs on three engagements and the design is solid, moving to "
    "Blueprint A with confidence is significantly de-risked."
)
challenge(
    "Is the staged investment argument actually credible, or are we just rationalizing Blueprint B because it "
    "is cheaper and easier? If we start with B and it works, what is the realistic probability we actually "
    "invest in rebuilding as A? Organizational momentum suggests we will just keep extending B indefinitely. "
    "Shawn: is there a pattern from other SNow partner practices that started with table extends and successfully "
    "migrated to a scoped app?"
)

# ── 9. ASSUMPTIONS TO VALIDATE ──────────────────────────────────────────────
doc.h1("Assumptions We Made That Should Be Validated")
doc.para(
    "This section is a candid list of things we assumed rather than verified. Each should be confirmed before "
    "Phase 1 begins."
)
doc.table(
    headers=["Assumption", "Why We Made It", "How to Validate", "Risk if Wrong"],
    rows=[
        ["Customers will use the portal if we build it", "Assumed based on how customers ask for status updates today", "Show a prototype to 2–3 customer project owners and ask if they would use it", "Medium: portal effort (~30% of build) wasted if adoption is low"],
        ["App Engine SKU is available on customer instances", "Most ECS engagements target customers with mid-to-large SNow footprints", "Check current customer portfolio: what % have App Engine? Check next 3 prospects", "High for Blueprint A: if most customers lack App Engine, Blueprint B becomes the only viable option"],
        ["Performance Analytics (PA) is available on customer instances", "PA is commonly licensed with ITSM Pro and above", "Check same customer portfolio for PA licensing", "Low: we designed native reporting fallback into both blueprints"],
        ["Existing ECS collateral can be loaded as methodology content without re-authoring", "All guides, pre-reads, cheatsheets, and accelerator packs already exist", "Content Owner to audit all 18 accelerator packs + guides for KB-tagging readiness", "Low: content exists; tagging is effort, not authoring"],
        ["One dedicated developer can build this solo", "Both blueprints were scoped to a single primary developer to keep the model simple", "Check availability of developer resource; validate scope estimate with a senior SNow developer", "High: if the developer is shared across delivery, Phase 1–3 timelines extend significantly"],
        ["The health score formula is correct", "Designed by delivery judgment, not historical data", "Run the formula retrospectively against 2–3 completed engagements to see if RAG matches actual outcomes", "Medium: a wrong formula gives customers a misleading health signal — worse than no formula"],
    ],
    col_widths_in=[2.0, 2.0, 2.5, 2.86],
)

# ── 10. OPEN QUESTIONS FOR THE REVIEW ───────────────────────────────────────
doc.h1("Open Questions for the Architecture Review")
doc.para("These are the questions the Shawn review session should resolve:")
doc.table(
    headers=["Topic", "Question", "Owner"],
    rows=[
        ["Architecture", "Blueprint A or Blueprint B — which do we start with?", "Shawn + Senior Director"],
        ["Developer resource", "Is a dedicated developer available? If not, what is the realistic calendar impact of a shared resource?", "Shawn + Senior Director"],
        ["Store commitment", "Is the ServiceNow Store a real near-term goal, or is it aspirational? If real, Blueprint A is mandatory. If aspirational, Blueprint B may be sufficient.", "Shawn + Senior Director"],
        ["Content table", "Should the Methodology Content table be its own custom table (Blueprint A) or extend KB (Blueprint B approach)? Worth reconsidering even in a Blueprint A context.", "Shawn + Senior Director"],
        ["Health formula validation", "Can we validate the health score formula against historical engagement data before building it in?", "Shawn + Senior Director"],
        ["Customer research", "Is there a willing customer or prospect we can prototype the portal with before committing to the Phase 2 build?", "Shawn + Senior Director"],
        ["Store competitive scan", "Has a thorough ServiceNow Store competitive scan been done? Should be done before Phase 0.", "Shawn + Senior Director"],
        ["App Engine availability", "What % of our current and pipeline customers have App Engine? This materially affects the Blueprint A/B decision.", "Shawn + Senior Director"],
        ["Post-launch ops", "How do we handle app updates across all customer instances post-launch? This is not designed yet.", "Shawn + Senior Director"],
        ["Blueprint B exit trigger", "If we start with Blueprint B and it works, what is the specific trigger that makes us invest in Blueprint A? Without a defined trigger, Blueprint B becomes permanent by default.", "Shawn + Senior Director"],
    ],
    col_widths_in=[2.0, 5.56, 1.8],
)

doc.save(OUT)
print(f"Saved: {OUT}")
