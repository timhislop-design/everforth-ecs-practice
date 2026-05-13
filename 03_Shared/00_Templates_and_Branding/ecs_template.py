"""
ecs_template.py — Everforth ECS Federal canonical document builder

The single source of truth for ECS Internal collateral formatting. Every script-built
document in the practice library imports from this module so that branding, typography,
table styling, headers, footers, and cover-page structure stay identical across artifacts.

Brand reference: ECS_Internal_Governance_Operating_Guide.docx (01_Internal/05_Discipline_How-To_Guides/)

Usage example
-------------
    from ecs_template import EcsDocument

    doc = EcsDocument(
        meta=dict(
            eyebrow="INTERNAL MANAGEMENT PLAYBOOK",
            title="Manager's Trust-But-Verify\\nPlaybook",
            subtitle="How ECS practice management catches OOTB-first drift",
            org="ECS Federal · ServiceNow Practice",
            audience="Practice Lead, Engagement Managers, Solution Architects",
            companion_to="Internal Governance Operating Guide · OOTB Delivery Playbook",
            doc_id="INT-TBV-01",
            version="1.0",
            status="Released",
            confidentiality="Internal Use Only · Confidential",
            running_header_label="Internal · Manager's Trust-But-Verify Playbook",
        ),
    )
    doc.add_cover_page()
    doc.add_page_break()
    doc.h1("How to Use This Playbook", numbered=False)
    doc.para("…")
    doc.h1("The Trust-But-Verify Discipline", numbered=True)  # auto-numbers as "1."
    doc.h2("Why the discipline exists")
    doc.table(headers=[…], rows=[…])
    doc.save("output.docx")
"""

from __future__ import annotations
from copy import deepcopy
import os
import re
import shutil
import zipfile
from dataclasses import dataclass, field
from typing import List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# BRAND CONSTANTS — DO NOT CHANGE WITHOUT PRACTICE LEAD APPROVAL
# =============================================================================
# Colors are stored as both hex strings (for XML attrs) and RGBColor (for python-docx APIs).

class Brand:
    NAVY_HEX        = "0B1F3A"   # Primary heading + table header fill
    TEAL_BRIGHT_HEX = "14B8A6"   # Eyebrow tag + divider lines
    TEAL_DEEP_HEX   = "0D9488"   # Subtitle italic
    ACCENT_BLUE_HEX = "2E74B5"   # H3 sub-sections
    SLATE_HEX       = "475569"   # Header right-side label, footer, audience meta
    BODY_HEX        = "1A1A1A"   # Body text dark gray
    ALT_ROW_HEX     = "F8FAFC"   # Table alternating row shading
    BORDER_HEX      = "E2E8F0"   # Table cell borders
    CALLOUT_BG_HEX  = "EAF1F8"   # Callout box fill
    WHITE_HEX       = "FFFFFF"

    NAVY        = RGBColor(0x0B, 0x1F, 0x3A)
    TEAL_BRIGHT = RGBColor(0x14, 0xB8, 0xA6)
    TEAL_DEEP   = RGBColor(0x0D, 0x94, 0x88)
    ACCENT_BLUE = RGBColor(0x2E, 0x74, 0xB5)
    SLATE       = RGBColor(0x47, 0x55, 0x69)
    BODY        = RGBColor(0x1A, 0x1A, 0x1A)
    WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


# Default logo path — repo-relative, assumes scripts run from anywhere in the repo.
LOGO_DEFAULT = "/sessions/sharp-magical-feynman/mnt/everforth-ecs-practice/00_Master_Blueprint/assets/everforth_logo.png"


# =============================================================================
# OOXML schema-order helpers (python-docx does not enforce these)
# =============================================================================
PPR_ORDER = [
    'pStyle','keepNext','keepLines','pageBreakBefore','framePr','widowControl',
    'numPr','suppressLineNumbers','pBdr','shd','tabs','suppressAutoHyphens',
    'kinsoku','wordWrap','overflowPunct','topLinePunct','autoSpaceDE','autoSpaceDN',
    'bidi','adjustRightInd','snapToGrid','spacing','ind','contextualSpacing',
    'mirrorIndents','suppressOverlap','jc','textDirection','textAlignment',
    'textboxTightWrap','outlineLvl','divId','cnfStyle','rPr','sectPr','pPrChange'
]
TCPR_ORDER = [
    'cnfStyle','tcW','gridSpan','hMerge','vMerge','tcBorders','shd','noWrap',
    'tcMar','textDirection','tcFitText','vAlign','hideMark','headers',
    'cellIns','cellDel','cellMerge','tcPrChange'
]
RPR_ORDER = [
    'rStyle','rFonts','b','bCs','i','iCs','caps','smallCaps','strike','dstrike',
    'outline','shadow','emboss','imprint','noProof','snapToGrid','vanish','webHidden',
    'color','spacing','w','kern','position','sz','szCs','highlight','u','effect',
    'bdr','shd','fitText','vertAlign','rtl','cs','em','lang','eastAsianLayout','specVanish','oMath'
]


def _local(elem) -> str:
    return elem.tag.split('}')[-1]


def insert_in_order(parent, new_elem, order_list):
    """Insert new_elem at the schema-correct position within parent."""
    new_local = _local(new_elem)
    if new_local not in order_list:
        parent.append(new_elem)
        return
    new_idx = order_list.index(new_local)
    insert_before = None
    for child in list(parent):
        local = _local(child)
        if local in order_list and order_list.index(local) > new_idx:
            insert_before = child
            break
    if insert_before is None:
        parent.append(new_elem)
    else:
        parent.insert(list(parent).index(insert_before), new_elem)


# =============================================================================
# Document metadata
# =============================================================================
@dataclass
class DocMeta:
    eyebrow: str                          # e.g. "INTERNAL MANAGEMENT PLAYBOOK"
    title: str                            # main title (use \n for soft breaks)
    subtitle: str = ""                    # italic deep-teal line under title
    org: str = "ECS Federal · ServiceNow Practice"
    audience: str = ""                    # rendered as "Audience: …"
    companion_to: str = ""                # rendered as "Companion to: …"
    doc_id: str = ""                      # e.g. "INT-TBV-01"
    version: str = "1.0"
    status: str = "Released"
    confidentiality: str = "Internal Use Only · Confidential"
    running_header_label: str = ""        # right-side header text on body pages
    footer_left: str = ""                 # OPTIONAL — overrides the default footer-left text.
                                          # Default (when empty): "{org} · Internal Use Only" — for Internal.
                                          # Client artifacts should pass: "{org} · Confidential" or similar non-internal phrasing.


# =============================================================================
# Main builder class
# =============================================================================
class EcsDocument:
    def __init__(self, meta: DocMeta | dict, logo_path: str = LOGO_DEFAULT):
        if isinstance(meta, dict):
            meta = DocMeta(**meta)
        self.meta = meta
        self.logo_path = logo_path
        self.doc = Document()
        self._h1_counter = 0   # auto section numbering
        self._configure_page()
        self._configure_styles()
        self._configure_header_footer()

    # --------- one-time setup ---------
    def _configure_page(self):
        s = self.doc.sections[0]
        s.page_width = Inches(8.5)
        s.page_height = Inches(11)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.different_first_page_header_footer = False
        s.header_distance = Inches(0.35)
        s.footer_distance = Inches(0.4)

    def _force_calibri(self, rpr_elem):
        rfonts = rpr_elem.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            insert_in_order(rpr_elem, rfonts, RPR_ORDER)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rfonts.set(qn(attr), 'Calibri')

    def _configure_styles(self):
        d = self.doc

        # Default body
        normal = d.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.font.color.rgb = Brand.BODY
        self._force_calibri(normal.element.get_or_add_rPr())

        def style_heading(name, size_pt, color_rgb, before=12, after=6):
            s = d.styles[name]
            s.font.name = "Calibri"
            s.font.size = Pt(size_pt)
            s.font.bold = True
            s.font.color.rgb = color_rgb
            pf = s.paragraph_format
            pf.space_before = Pt(before)
            pf.space_after = Pt(after)
            pf.keep_with_next = True
            self._force_calibri(s.element.get_or_add_rPr())

        style_heading("Heading 1", 16, Brand.NAVY,        before=18, after=6)
        style_heading("Heading 2", 13, Brand.NAVY,        before=14, after=4)
        style_heading("Heading 3", 11.5, Brand.ACCENT_BLUE, before=10, after=2)
        style_heading("Title",     28, Brand.NAVY,        before=0,  after=4)
        # Subtitle is italic deep teal
        sub = d.styles["Subtitle"] if "Subtitle" in [s.name for s in d.styles] else None
        # python-docx may not expose Subtitle by name in all versions; we just
        # apply formatting inline at the cover_page step.

    def _configure_header_footer(self):
        section = self.doc.sections[0]
        # ---- HEADER ----
        header = section.header
        p = header.paragraphs[0]
        # Tab stops: right tab at end of content area
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9360')
        tabs.append(tab)
        existing = pPr.find(qn('w:tabs'))
        if existing is not None:
            pPr.remove(existing)
        insert_in_order(pPr, tabs, PPR_ORDER)
        # Bottom border (teal divider line under header)
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:color'), Brand.TEAL_BRIGHT_HEX)
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        insert_in_order(pPr, pBdr, PPR_ORDER)
        # Logo — use header part's new_pic_inline directly; run.add_picture() fails
        # in header context on some python-docx builds (IndexError in part resolution)
        logo_run = p.add_run()
        if os.path.exists(self.logo_path):
            try:
                from PIL import Image
                import io as _io
                _img = Image.open(self.logo_path).convert("RGBA")
                _buf = _io.BytesIO()
                _img.save(_buf, format="PNG")
                _buf.seek(0)
                _hpart = header.part
                _inline = _hpart.new_pic_inline(_buf, Inches(1.05), None)
                logo_run._r.add_drawing(_inline)
            except Exception:
                pass  # logo silently absent rather than crashing the build
        # Tab to right
        p.add_run("\t")
        # Right label (letter-spaced slate)
        if self.meta.running_header_label:
            r = p.add_run(self.meta.running_header_label)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            r.font.color.rgb = Brand.SLATE
            # letter spacing
            rPr = r._r.get_or_add_rPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:val'), '30')
            insert_in_order(rPr, spacing, RPR_ORDER)

        # ---- FOOTER ----
        footer = section.footer
        fp = footer.paragraphs[0]
        pPr2 = fp._p.get_or_add_pPr()
        tabs2 = OxmlElement('w:tabs')
        tab2 = OxmlElement('w:tab')
        tab2.set(qn('w:val'), 'right')
        tab2.set(qn('w:pos'), '9360')
        tabs2.append(tab2)
        existing2 = pPr2.find(qn('w:tabs'))
        if existing2 is not None:
            pPr2.remove(existing2)
        insert_in_order(pPr2, tabs2, PPR_ORDER)
        # Left text
        footer_left_text = self.meta.footer_left or f"{self.meta.org}  ·  Internal Use Only"
        r1 = fp.add_run(footer_left_text)
        r1.font.name = "Calibri"; r1.font.size = Pt(8); r1.font.color.rgb = Brand.SLATE
        # Tab + right page-number
        fp.add_run("\t")
        r2 = fp.add_run("Page ")
        r2.font.name = "Calibri"; r2.font.size = Pt(8); r2.font.color.rgb = Brand.SLATE
        self._add_field(fp, "PAGE")
        r3 = fp.add_run(" of ")
        r3.font.name = "Calibri"; r3.font.size = Pt(8); r3.font.color.rgb = Brand.SLATE
        self._add_field(fp, "NUMPAGES")

    def _add_field(self, paragraph, code: str):
        """Append a Word field (e.g., PAGE, NUMPAGES) to a paragraph."""
        run = paragraph.add_run()
        run.font.name = "Calibri"; run.font.size = Pt(8); run.font.color.rgb = Brand.SLATE
        for kind, val in [('begin', None), ('instr', code), ('separate', None), ('end', None)]:
            if kind == 'instr':
                e = OxmlElement('w:instrText'); e.text = val
            else:
                e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), kind)
            run._r.append(e)

    # =========================================================================
    # CONTENT PRIMITIVES
    # =========================================================================

    def add_cover_page(self):
        """Render the canonical cover page block. Caller adds add_page_break() after."""
        m = self.meta
        # Vertical breathing space (top of page)
        for _ in range(2):
            sp = self.doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(6)
        # Eyebrow tag (BRIGHT TEAL bold uppercase letter-spaced)
        if m.eyebrow:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(40)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(m.eyebrow.upper())
            r.font.name = "Calibri"; r.font.size = Pt(9); r.font.bold = True
            r.font.color.rgb = Brand.TEAL_BRIGHT
            rPr = r._r.get_or_add_rPr()
            sp_e = OxmlElement('w:spacing'); sp_e.set(qn('w:val'), '60')
            insert_in_order(rPr, sp_e, RPR_ORDER)
        # Title (NAVY bold, multi-line preserved)
        for line in m.title.split("\n"):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            r.font.name = "Calibri"; r.font.size = Pt(28); r.font.bold = True
            r.font.color.rgb = Brand.NAVY
        # Subtitle (italic deep teal)
        if m.subtitle:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(20)
            r = p.add_run(m.subtitle)
            r.font.name = "Calibri"; r.font.size = Pt(12); r.font.italic = True
            r.font.color.rgb = Brand.TEAL_DEEP
        # Divider line (paragraph border bottom, BRIGHT TEAL)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), Brand.TEAL_BRIGHT_HEX)
        pBdr.append(bottom)
        insert_in_order(pPr, pBdr, PPR_ORDER)
        # Org, audience, companion-to, doc id, confidentiality
        def line(text, italic=False, bold=False, color=Brand.BODY, size=11, after=2):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(after)
            r = p.add_run(text)
            r.font.name = "Calibri"; r.font.size = Pt(size)
            r.font.bold = bold; r.font.italic = italic
            r.font.color.rgb = color
            return p
        line(m.org, bold=True, color=Brand.NAVY, size=11, after=2)
        if m.audience:
            line(f"Audience: {m.audience}", color=Brand.SLATE, size=10, after=2)
        if m.companion_to:
            line(f"Companion to: {m.companion_to}", color=Brand.SLATE, size=10, after=2)
        meta_bits = []
        if m.doc_id:   meta_bits.append(f"Document ID: {m.doc_id}")
        if m.version:  meta_bits.append(f"Version: {m.version}")
        if m.status:   meta_bits.append(f"Status: {m.status}")
        if meta_bits:
            line("     ·     ".join(meta_bits), color=Brand.SLATE, size=10, after=4)
        if m.confidentiality:
            line(m.confidentiality, italic=True, color=Brand.SLATE, size=10, after=14)

    # ---- structural ----
    def page_break(self):
        p = self.doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    def add_page_break(self):  # alias
        self.page_break()

    # ---- headings ----
    def h1(self, text: str, numbered: bool = True):
        if numbered:
            self._h1_counter += 1
            text = f"{self._h1_counter}.  {text}"
        return self.doc.add_paragraph(text, style="Heading 1")

    def h2(self, text: str):
        return self.doc.add_paragraph(text, style="Heading 2")

    def h3(self, text: str):
        return self.doc.add_paragraph(text, style="Heading 3")

    # ---- prose ----
    def para(self, text: str, *, italic=False, bold=False, color=None,
             size=11, space_after=8, space_before=0, align=None):
        p = self.doc.add_paragraph()
        if align: p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(size)
        if bold:   r.font.bold = True
        if italic: r.font.italic = True
        r.font.color.rgb = color or Brand.BODY
        return p

    def bullet(self, text: str, *, level: int = 0):
        p = self.doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
        p.paragraph_format.space_after = Pt(2)
        if p.runs:
            r = p.runs[0]; r.text = text
        else:
            r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(11)
        return p

    def callout(self, text: str, *, color_rgb=None):
        color_rgb = color_rgb or Brand.NAVY
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        # Border
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '4'); b.set(qn('w:color'), Brand.ACCENT_BLUE_HEX)
            pBdr.append(b)
        insert_in_order(pPr, pBdr, PPR_ORDER)
        # Shading
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), Brand.CALLOUT_BG_HEX)
        insert_in_order(pPr, shd, PPR_ORDER)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = color_rgb
        return p

    # ---- table (canonical: navy header, alternating row shading) ----
    def table(self, *, headers: List[str], rows: List[List[str]],
              col_widths_in: Optional[List[float]] = None,
              alt_shading: bool = True):
        n_cols = len(headers)
        if col_widths_in is None:
            # Equal widths summing to content width (9360 dxa)
            col_widths_in = [9.36 / n_cols] * n_cols
        t = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Apply column widths
        for i, w in enumerate(col_widths_in):
            for row in t.rows:
                row.cells[i].width = Inches(w)
        # HEADER ROW: navy fill, white bold text, mark as repeating
        hdr_row = t.rows[0]
        # Mark as table header (repeats on page break)
        trPr = hdr_row._tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            self._cell_shading(cell, Brand.NAVY_HEX)
            self._cell_borders(cell, Brand.BORDER_HEX)
            self._cell_margins(cell, top=120, left=160, bottom=120, right=160)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(h)
            r.font.name = "Calibri"; r.font.bold = True; r.font.size = Pt(11)
            r.font.color.rgb = Brand.WHITE
        # BODY ROWS — alternating shading
        for r_idx, row in enumerate(rows):
            shade = alt_shading and (r_idx % 2 == 0)
            for c_idx, val in enumerate(row):
                cell = t.rows[r_idx + 1].cells[c_idx]
                self._cell_borders(cell, Brand.BORDER_HEX)
                self._cell_margins(cell, top=100, left=160, bottom=100, right=160)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                if shade:
                    self._cell_shading(cell, Brand.ALT_ROW_HEX)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
                r = p.add_run(val)
                r.font.name = "Calibri"; r.font.size = Pt(10)
        # Trailing breath
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # ---- low-level cell helpers (public so build scripts can extend) ----
    def _cell_shading(self, cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')):
            tcPr.remove(old)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex)
        insert_in_order(tcPr, shd, TCPR_ORDER)

    def _cell_borders(self, cell, color_hex="E2E8F0"):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
            b.set(qn('w:color'), color_hex)
            tcBorders.append(b)
        insert_in_order(tcPr, tcBorders, TCPR_ORDER)

    def _cell_margins(self, cell, *, top=80, left=120, bottom=80, right=120):
        tcPr = cell._tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(old)
        tcMar = OxmlElement('w:tcMar')
        for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
            e = OxmlElement(f'w:{side}'); e.set(qn('w:type'), 'dxa'); e.set(qn('w:w'), str(val))
            tcMar.append(e)
        insert_in_order(tcPr, tcMar, TCPR_ORDER)

    # =========================================================================
    # SAVE — writes file, then patches OOXML schema-validation issues
    # =========================================================================
    def save(self, path: str):
        self.doc.save(path)
        self._patch_settings_xml(path)
        return path

    @staticmethod
    def _patch_settings_xml(path: str):
        """python-docx writes <w:zoom w:val="bestFit"/> with no percent attribute,
        which fails OOXML schema validation. Add w:percent="100"."""
        tmp = path + ".tmp"
        with zipfile.ZipFile(path, 'r') as zin:
            with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    data = zin.read(item)
                    if item == 'word/settings.xml':
                        txt = data.decode('utf-8')
                        txt = re.sub(
                            r'<w:zoom([^>]*)/>',
                            lambda m: m.group(0) if 'w:percent' in m.group(1)
                                      else f'<w:zoom{m.group(1)} w:percent="100"/>',
                            txt
                        )
                        data = txt.encode('utf-8')
                    zout.writestr(item, data)
        shutil.move(tmp, path)


# =============================================================================
# Convenience entry point — minimal example callable
# =============================================================================
if __name__ == "__main__":
    # Sanity-build a tiny doc to confirm the module imports cleanly
    d = EcsDocument(meta=DocMeta(
        eyebrow="INTERNAL — TEMPLATE TEST",
        title="ECS Template\nSmoke Test",
        subtitle="Confirms the canonical brand renders end-to-end",
        audience="Practice Lead",
        running_header_label="Internal · Template Smoke Test",
        doc_id="TEST-01",
    ))
    d.add_cover_page()
    d.page_break()
    d.h1("Headings render correctly")
    d.h2("Subheading")
    d.para("Body paragraph.")
    d.bullet("Bullet item")
    d.h3("Sub-subhead")
    d.table(headers=["Col A", "Col B", "Col C"],
            rows=[["a","b","c"], ["d","e","f"], ["g","h","i"]])
    d.callout("This is a callout.")
    out = "/tmp/ecs_template_smoke.docx"
    d.save(out)
    print(f"Smoke build saved: {out}")
